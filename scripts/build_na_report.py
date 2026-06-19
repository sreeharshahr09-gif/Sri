"""
Generate a 1-2 page Word report on the Nylon-Aramid hybrid-cord patents
(design / reverse-engineering / performance intelligence).

Numbers derive from the audited cohort in cohorts.py (single source of truth).
Output: reports/Nylon_Aramid_Design_Report.docx
Run   : python scripts/build_na_report.py
"""
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cohorts as C

# ── data ────────────────────────────────────────────────────────────────────
df = C.load(); m = C.masks(df); t = df["_txt"]
NA, WEAR = m["nylon_aramid"], m["wear"]
txt = t[NA.values]; sub = df[NA.values].copy()
N = int(NA.sum()); NWEAR = int((NA & WEAR).sum())
alive = sub["legal_status"].str.upper().eq("ALIVE")
N_ALIVE = int(alive.sum()); N_DEAD = N - N_ALIVE


def pct(pat):
    return round(100 * txt.str.contains(pat, regex=True).mean())


def norm(s):
    s = str(s).split(";")[0]; s = re.sub(r"\([A-Z]{2}\)", "", s)
    return re.sub(r"\s+", " ", s).strip(" .,").upper()


top_asg = sub["assignee"].map(norm).replace("", pd.NA).dropna().value_counts().head(6)
co_twist = int(txt.str.contains(r"co-?twist|cabled|twisted together|ply.?twist|folded together", regex=True).sum())

# ── document styling ────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x3B, 0x57)
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.8)


def H(text, size=15, space_before=8):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = NAVY
    p.space_before = Pt(space_before); p.space_after = Pt(3)
    return p


def body(text, bullet=False, after=3):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(after)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            p.add_run(seg[2:-2]).bold = True
        else:
            p.add_run(seg)
    return p


# ── title ───────────────────────────────────────────────────────────────────
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = ttl.add_run("Nylon-Aramid Hybrid Cords — Design, Reverse-Engineering & Performance")
r.bold = True; r.font.size = Pt(17); r.font.color.rgb = NAVY
sub_t = doc.add_paragraph()
r = sub_t.add_run(f"Zero-degree belt / cap-ply cords in PCR tires · {N} strict-cohort patents "
                  f"({NWEAR} with a wear claim) · audited, motorcycle-excluded")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── 1. summary ──────────────────────────────────────────────────────────────
H("1.  What this cord is, and why")
body("The Nylon-Aramid hybrid is a zero-degree cap-ply cord that pairs a low-shrink, "
     "high-modulus aramid filament with a flexible nylon-6.6 filament to give a two-stage "
     "(“J-shaped”) load–elongation response. Nylon carries the low-load region (flexibility, "
     "fatigue life, and post-cure shrink that tightens the cap ply during cure); aramid carries the "
     "high-load region (crown stiffness retained at speed and temperature). The design intent is "
     "high-speed durability and crown control — wear/mileage is a secondary, indirect benefit via "
     "footprint uniformity, not an abrasion-chemistry effect.")

# ── 2. construction recipe ──────────────────────────────────────────────────
H("2.  Canonical construction (reverse-engineering datasheet)")
body("Construction is consistent across assignees. A typical teardown shows:")
tbl = doc.add_table(rows=1, cols=2); tbl.style = "Light Grid Accent 1"
tbl.autofit = True
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("Parameter").bold = True
hdr[1].paragraphs[0].add_run("Typical value(s) in the corpus").bold = True
rows = [
    ("Aramid linear density", "1670 / 1400 / 1100 dtex (load-bearing filament)"),
    ("Nylon linear density", "940 dtex is the signature value; nylon-6.6 in ~half the cohort"),
    ("Ply / construction code", "mostly 2-ply — e.g. 1100/2, 1670/2, 1840/2 (some single-end)"),
    ("Twist", "200–500 TPM (peaks at 200, 300, 500), balanced ply-twist"),
    ("Cord density", "50–90 ends/dm (EPDM), centred ~80, spiral-wound cap"),
    ("Dual-modulus step", "modulus ratio ~1.5 — the defining two-stage signature"),
    ("Assembly style", f"co-twisted / cabled together dominates ({co_twist} patents); true core-sheath "
                       "is the minority — and is nylon-core / aramid-sheath more often than the reverse"),
]
for a, b in rows:
    c = tbl.add_row().cells
    c[0].paragraphs[0].add_run(a).bold = True
    c[1].paragraphs[0].add_run(b)
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(9.5)
body("")
body(f"Reverse-engineering tells: aramid is almost never branded ({pct(r'kevlar|twaron|technora')}% name "
     "Kevlar/Twaron/Technora) — a generic, low-heat-shrink high-modulus filament co-twisted with a "
     "940-dtex PA66 filament at ~300 TPM is the textbook hybrid. Dimensional-stability / heat-set terms "
     f"appear in {pct(r'shrinkage|heat.?set|dimensional stab|thermal')}% of patents (shrinkage ~0.5–4%); "
     "the aramid's low heat-shrink and the nylon's high shrink are the fingerprints in a cross-section.",
     bullet=True)

# ── 3. performance & failure mode ───────────────────────────────────────────
H("3.  Performance levers and the failure mode to manage")
body("Distinctive wear levers vs. an all-PCR baseline are in the cord itself: core-sheath/wrap (2.0×), "
     "dual-modulus (1.9×) and twist tuning (1.8×). Adhesion/dip and zoned density are NOT distinctive "
     "to Nylon-Aramid.", bullet=True)
body("The aramid stiffness step creates an uneven-wear risk at the cord's modulus transition. This is why "
     "the cohort leans defensive (fixing uneven wear) and why twist tuning and modulus grading recur — "
     "they exist largely to engineer that transition out. Pre-empt this failure mode in any new design.",
     bullet=True)

# ── 4. IP landscape ─────────────────────────────────────────────────────────
H("4.  IP landscape — design freedom & who to read")
body(f"Legal status: {N_DEAD} DEAD / {N_ALIVE} ALIVE. The large expired body is effectively a "
     "free-to-practice reverse-engineering library — the canonical recipe above is largely off-patent.",
     bullet=True)
body(f"Live art ({N_ALIVE} patents, {int((WEAR[sub.index] & alive).sum())} with a wear claim) clusters "
     "2018–2025 and is the FTO-relevant set — recent holders include Goodyear, Yokohama, Sumitomo, "
     "Pirelli and Hyosung. The wear-specific angle is thin and mostly dead: an opportunity, but validate "
     "against live PET/PA “center-vs-edge” zoned-density art.", bullet=True)
body("Where the know-how sits: " + ", ".join(f"{a.title()} ({n})" for a, n in top_asg.items())
     + ". For the yarn/cord engineering specifically, read the cord-makers Kordsa and Hyosung.", bullet=True)

# ── 5. takeaways ────────────────────────────────────────────────────────────
H("5.  Practical takeaways")
body("Specifying a cord: start from aramid ~1670 dtex + nylon-6.6 940 dtex, 2-ply, ~300 TPM, ~80 ends/dm "
     "spiral cap, ~1.5× modulus step; then tune twist/grading to suppress the uneven-wear transition.",
     bullet=True)
body("Reverse-engineering a competitor: the co-twist vs. nylon-core/aramid-sheath distinction, the "
     "940-dtex nylon tell, and the aramid low-heat-shrink signature are your fingerprints.", bullet=True)
body("Positioning performance: sell high-speed / high-temperature durability and crown stiffness; treat "
     "mileage as a derived benefit; pre-empt the uneven-wear failure mode.", bullet=True)

# ── footer note ─────────────────────────────────────────────────────────────
note = doc.add_paragraph()
r = note.add_run("Method note: figures are text-mined from PatSeer fields + AI summaries over the strict "
                 f"Nylon-Aramid cohort (n={N}); high recall — spot-check claims before quoting an "
                 "individual patent. Corpus skews pre-2015; legal status as of the export date.")
r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

out = "reports/Nylon_Aramid_Design_Report.docx"
doc.save(out)
print("saved", out, "| paragraphs", len(doc.paragraphs))
